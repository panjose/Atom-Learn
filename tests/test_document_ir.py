from __future__ import annotations

import hashlib

import pytest

import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "atom-learn" / "scripts"))

from document_ir import DocumentIRError, build_document_ir
from rag import extract_path


def test_table_ir_preserves_headers_spans_and_review_metadata() -> None:
    document = build_document_ir(
        source_id="paper.tables",
        source_revision=1,
        title="Paper",
        uri="inline:paper.tables",
        suffix=".md",
        sections=[
            {
                "locator": "page 2, table 1",
                "section": "Results",
                "text": "| Metric | Value |\n| --- | --- |\n| Accuracy | 0.91 |",
                "kind": "table",
            }
        ],
    )
    table = next(block for block in document["blocks"] if block["kind"] == "table")
    assert table["page"] == 2
    assert table["review_status"] == "reviewed"
    assert table["table_structure"]["header_rows"] == 1
    assert table["table_structure"]["rows"][0]["cells"][0]["is_header"] is True
    assert table["table_structure"]["rows"][1]["cells"][1]["text"] == "0.91"


def test_vision_figure_numeric_values_are_proposals_until_reviewed() -> None:
    crop_hash = "sha256:" + hashlib.sha256(b"figure-crop").hexdigest()
    document = build_document_ir(
        source_id="paper.figure",
        source_revision=1,
        title="Paper",
        uri="inline:paper.figure",
        suffix=".json",
        sections=[
            {
                "locator": "page 4, figure 2",
                "section": "Results",
                "text": "Figure 2: treatment effect with error bars",
                "kind": "figure",
                "bbox": [10, 20, 300, 400],
                "crop_hash": crop_hash,
                "extraction_method": "harness_vision",
            }
        ],
    )
    figure = next(block for block in document["blocks"] if block["kind"] == "figure")
    assert figure["crop_hash"] == crop_hash
    assert figure["review_status"] == "proposed"
    assert figure["numeric_status"] == "proposal"


def test_visual_blocks_retain_adjacent_prose_locators() -> None:
    document = build_document_ir(
        source_id="paper.adjacent",
        source_revision=1,
        title="Paper",
        uri="inline:paper.adjacent",
        suffix=".json",
        sections=[
            {"locator": "page 3, before", "section": "Results", "text": "The result is discussed here."},
            {"locator": "page 3, figure 1", "section": "Results", "text": "Figure 1", "kind": "figure"},
            {"locator": "page 3, after", "section": "Results", "text": "The caption explains the comparison."},
        ],
    )
    figure = next(block for block in document["blocks"] if block["kind"] == "figure")
    assert len(figure["adjacent_block_ids"]) == 2


def test_ir_rejects_unknown_caption_reference() -> None:
    with pytest.raises(DocumentIRError, match="missing caption_block_id"):
        build_document_ir(
            source_id="paper.bad-caption",
            source_revision=1,
            title="Paper",
            uri="inline:paper.bad-caption",
            suffix=".json",
            sections=[
                {
                    "locator": "page 1, figure 1",
                    "section": "Results",
                    "text": "Figure",
                    "kind": "figure",
                    "caption_block_id": "block-000000000000000000000000",
                }
            ],
        )


def test_html_extractor_preserves_header_and_column_spans(tmp_path: Path) -> None:
    source = tmp_path / "table.html"
    source.write_text(
        "<table><tr><th colspan='2'>Results</th></tr><tr><td>A</td><td>0.9</td></tr></table>",
        encoding="utf-8",
    )
    document = build_document_ir(
        source_id="paper.html-table",
        source_revision=1,
        title="Paper",
        uri=str(source),
        suffix=".html",
        sections=extract_path(source),
    )
    table = next(block for block in document["blocks"] if block["kind"] == "table")
    assert table["table_structure"]["header_rows"] == 1
    assert table["table_structure"]["rows"][0]["cells"][0]["column_span"] == 2


def test_docx_extractor_preserves_merged_cell_spans(tmp_path: Path) -> None:
    from docx import Document

    source = tmp_path / "table.docx"
    document = Document()
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).merge(table.cell(0, 1)).text = "Results"
    table.cell(1, 0).text = "A"
    table.cell(1, 1).text = "0.9"
    document.save(source)
    ir = build_document_ir(
        source_id="paper.docx-table",
        source_revision=1,
        title="Paper",
        uri=str(source),
        suffix=".docx",
        sections=extract_path(source),
    )
    ir_table = next(block for block in ir["blocks"] if block["kind"] == "table")
    assert ir_table["table_structure"]["rows"][0]["cells"][0]["column_span"] == 2
